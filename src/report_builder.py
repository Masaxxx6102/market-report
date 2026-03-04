from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
from typing import List, Dict, Any
from .logger_setup import logger

class ReportBuilder:
    """
    Constructs the professional Word (.docx) report.
    """
    OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "output")

    def __init__(self, organization: str = "Forcus株式会社"):
        self.organization = organization
        if not os.path.exists(self.OUTPUT_DIR):
            os.makedirs(self.OUTPUT_DIR)
            
    def build_report(self, data: Dict[str, Any], date_str: str = None) -> str:
        if not date_str:
            date_str = datetime.now().strftime("%Y-%m-%d")
        
        doc = Document()
        
        # --- Title ---
        title = doc.add_heading("Daily Market Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_para = doc.add_paragraph(f"日付: {date_str}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        org_para = doc.add_paragraph(f"作成: {self.organization}")
        org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # doc.add_page_break() # Remove page break to fill 1st page

        # --- Section 1: Charts ---
        doc.add_heading("1. 主要指数チャート", 1)
        chart_data = data.get("charts", {})
        for idx_name, paths in chart_data.items():
            doc.add_heading(idx_name, 2)
            # Use a table to place long and short charts side-by-side
            table = doc.add_table(rows=1, cols=2)
            table.autofit = True
            cells = table.rows[0].cells
            
            if "short" in paths and os.path.exists(paths["short"]):
                self._add_image_to_cell(cells[0], paths["short"], width=Inches(3.0))
            if "long" in paths and os.path.exists(paths["long"]):
                self._add_image_to_cell(cells[1], paths["long"], width=Inches(3.0))
            doc.add_paragraph() # Spacer

        # --- Section 2: Market Table ---
        doc.add_heading("2. 株価概況", 1)
        quote_data = data.get("quotes", {})
        if quote_data:
            # Group by Market/US/JP
            market_quotes = {k: v for k, v in quote_data.items() if "^" in k or k == "TOPIX"}
            us_quotes = {k: v for k, v in quote_data.items() if "^" not in k and k != "TOPIX" and ".T" not in k}
            jp_quotes = {k: v for k, v in quote_data.items() if ".T" in k}
            
            if market_quotes:
                doc.add_heading("【マーケット状況】", 2)
                # Sort: Nikkei 225 (^N225) first
                sorted_market = {}
                if "^N225" in market_quotes:
                    sorted_market["^N225"] = market_quotes["^N225"]
                for k, v in market_quotes.items():
                    if k != "^N225":
                        sorted_market[k] = v
                self._add_quote_table(doc, sorted_market)
            if us_quotes:
                doc.add_heading("【米国株 時価総額上位20】", 2)
                self._add_quote_table(doc, us_quotes)
            if jp_quotes:
                doc.add_heading("【日本株 時価総額上位20】", 2)
                self._add_quote_table(doc, jp_quotes)

        # --- Section 3: Market Overview ---
        overview = data.get("overview", {})
        if overview:
            doc.add_heading("3. マーケット概況", 1)
            if "jp_overview" in overview:
                doc.add_heading("【日本市場】", 2)
                p = doc.add_paragraph(overview["jp_overview"])
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if "us_overview" in overview:
                doc.add_heading("【米国市場】", 2)
                p = doc.add_paragraph(overview["us_overview"])
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # --- Section 4: News ---
        news = data.get("news", [])
        if news:
            doc.add_heading("4. 厳選ニュース", 1)
            
            # Group news by category
            categories = {}
            for item in news:
                cat = item.get("category", "【その他】")
                if cat not in categories:
                    categories[cat] = []
                categories[cat].append(item)
            
            # Define preferred order
            order = ["【マクロ経済・金融政策】", "【メガキャップ・ムーブメント】", "【決算・ガイダンス速報】", "【国際・地政学】"]
            
            for cat in order:
                if cat in categories:
                    doc.add_heading(cat, 2)
                    for item in categories[cat]:
                        self._add_news_item(doc, item)
            
            # Any other categories
            for cat, items in categories.items():
                if cat not in order:
                    doc.add_heading(cat, 2)
                    for item in items:
                        self._add_news_item(doc, item)

        # --- Section 5: Earnings ---
        earnings = data.get("earnings", [])
        if earnings:
            doc.add_heading("5. 決算・トピックス", 1)
            self._add_earnings_table(doc, earnings)

        # Save
        filename = f"daily-report-{date_str}.docx"
        output_path = os.path.join(self.OUTPUT_DIR, filename)
        doc.save(output_path)
        logger.info(f"Report saved to {output_path}")
        return output_path

    def _add_news_item(self, doc, item):
        p = doc.add_paragraph()
        importance = item.get("importance", 5)
        star = "★ " if importance >= 8 else ""
        run = p.add_run(f"{star}{item.get('headline')}")
        run.bold = True
        
        doc.add_paragraph(item.get("summary"), style='List Bullet')
        source_para = doc.add_paragraph(f"出典: {item.get('source')} | 関連: {', '.join(item.get('related_symbols', []))}")
        source_para.paragraph_format.left_indent = Inches(0.5)

    def _set_cell_style(self, cell, text, color=None, bold=False):
        para = cell.paragraphs[0]
        run = para.add_run(str(text))
        if color:
            run.font.color.rgb = color
        if bold:
            run.bold = True

    def _add_quote_table(self, doc, quotes: Dict[str, Dict[str, Any]]):
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        # Updated Headers: [Name, Ticker, Close, DailyChg, Daily%, YTDChg, YTD%]
        hdrs = ["銘柄名", "コード", "終値", "前日比", "前日比%", "年初来比", "年初来比%"]
        for i, h in enumerate(hdrs):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            
        for sym, q in quotes.items():
            row_cells = table.add_row().cells
            row_cells[0].text = q.get("name", "")[:20] 
            row_cells[1].text = sym
            row_cells[2].text = f"{q.get('price'):,.2f}"
            
            # Daily Change (Value)
            change = q.get("change", 0)
            color = RGBColor(0xCC, 0x00, 0x00) if change < 0 else RGBColor(0x00, 0x66, 0x00)
            self._set_cell_style(row_cells[3], f"{change:+,.2f}", color=color)
            
            # Daily %
            cp = q.get("changesPercentage", 0)
            self._set_cell_style(row_cells[4], f"{cp:+.2f}%", color=color)
            
            # YTD Change (Value)
            ychange = q.get("ytdChange", 0)
            ycolor = RGBColor(0xCC, 0x00, 0x00) if ychange < 0 else RGBColor(0x00, 0x66, 0x00)
            self._set_cell_style(row_cells[5], f"{ychange:+,.2f}", color=ycolor)
            
            # YTD %
            yp = q.get("ytdChangePercentage", 0)
            self._set_cell_style(row_cells[6], f"{yp:+.2f}%", color=ycolor)

    def _add_image_to_cell(self, cell, path, width):
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(path, width=width)

    def _add_earnings_table(self, doc, earnings: List[Dict[str, Any]]):
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdrs = ["銘柄", "EPS実績", "EPS予想", "サプライズ", "判定"]
        for i, h in enumerate(hdrs):
            hdr_cells[i].text = h
            
        for e in earnings:
            row_cells = table.add_row().cells
            row_cells[0].text = e["symbol"]
            row_cells[1].text = f"{e['reported_eps']:.2f}"
            row_cells[2].text = f"{e['consensus_eps']:.2f}"
            row_cells[3].text = f"{e['surprise_pct']:.1f}%"
            
            res = e["result"]
            color = RGBColor(0x00, 0x66, 0x00) if res == "Beat" else RGBColor(0xCC, 0x00, 0x00) if res == "Miss" else None
            self._set_cell_style(row_cells[4], res, color=color, bold=True)
